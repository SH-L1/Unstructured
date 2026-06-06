"""Generate concise trust-workflow reference documents.

The generated documents describe only the versioned V1 workflow. They never
present ranking diagnostics as confidence and never describe an automatic
send or completion path.
"""

from pathlib import Path

from docx import Document


OUT_DIR = Path("docs/corrected-documents")
DATE = "2026-06-04"
LEGACY_CORRECTED_REPLACEMENTS = {
    "API Key SHA-256 해시": "레거시 API 키 저장소 제거됨",
    "API Key 인증 사용자": "세션 역할 사용자",
    "API Key 인증 사용자와 해시 저장": "레거시 API 키 사용자·해시 저장소 제거",
    "RECEIVED, ANALYZED, DRAFT_GENERATED 등": (
        "RECEIVED, TRIAGE_REVIEW, DRAFT_REVIEW, APPROVAL_PENDING, APPROVED, COMPLETED"
    ),
    "API Key": "세션 로그인 및 CSRF",
    "민원 분석 Agent Loop 결과 화면": "근거 검증형 민원 검토 작업 공간",
    "민원 원문 입력, 채널/위치/API Key 입력, Agent Loop 진행 표시, 결과 패널": (
        "비식별 민원 접수, 채널·위치 확인, 작업 상태, 근거·검증·사람 결정 패널"
    ),
    "초안 수정 PUT /draft와 첨부파일 API는 현재 대시보드 편집 UI가 아니라 Swagger/API 통합시험 범위이다.": (
        "첨부·위치 확인·검토·승인·수동 완료 변경 API는 역할, Idempotency-Key, If-Match를 요구한다."
    ),
    "API Key 선택 활성화, audit 기본 활성화": "세션 역할 인증 필수, 변경 감사 기본 활성화",
    "API Key 인증과 감사 로그는 설정값으로 활성화되며 api_users/audit_logs를 사용한다.": (
        "세션 역할 인증과 감사 로그를 사용하며 레거시 api_users와 브라우저 API 키는 제거했다."
    ),
    "API Key 인증과 감사 추적을 지원해야 함": "세션 역할 인증과 모든 변경 API 감사 추적을 지원해야 함",
    "GET /analysis 호출 시 OpenAI/Bedrock/룰 기반 분석으로 complaint_analysis를 생성한다.": (
        "POST /api/v1/complaints/{id}/analysis-runs로 Python 분석 작업을 요청하고 "
        "Spring이 검증 결과를 complaint_analysis와 complaint_issues에 적용한다."
    ),
    "GET /draft 호출 시 지식문서 검색 후 official_drafts와 rag_contexts를 저장한다.": (
        "POST /api/v1/complaints/{id}/draft-runs로 근거 제한 초안을 요청하고 Spring이 "
        "주장·근거·스냅샷을 검증한 뒤 검토용 초안을 저장한다."
    ),
    "GET /rag-contexts와 화면 렌더링으로 근거 문서 목록을 확인한다.": (
        "GET /api/v1/complaints/{id}와 검토 화면에서 불변 근거 스냅샷과 검증 결과를 확인한다."
    ),
    "PUT /draft로 초안을 수정하면 draft_revisions에 변경 이력이 저장된다.": (
        "POST /api/v1/drafts/{id}/reviews와 /approvals로 분리된 사람 결정을 기록한다."
    ),
    "민원 원문 입력, Agent Loop 진행 상태, 분석 JSON, RAG 근거, 공문 초안 확인": (
        "비식별 민원 입력, 작업 상태, 구조화 분석, 근거 스냅샷, 검증 실패, 검토용 초안 확인"
    ),
    "API Key 인증, 감사 로그, Actuator, OpenAPI": (
        "세션 역할 인증, 감사 로그, 제한 Actuator, 운영 비공개 OpenAPI"
    ),
    "gpt-4o-mini 기반 분석/초안 생성, API Key 설정 시 사용": (
        "Python 작업자의 구성된 AI 공급자로 구조화 분석·초안 제안"
    ),
    "health, Swagger, API Key 인증, 감사 로그 확인": (
        "health 권한, 운영 Swagger 차단, 세션 역할 인증, 감사 로그 확인"
    ),
    "Swagger, health, API Key, 감사 로그, 테스트 확인": (
        "운영 Swagger 차단, health 권한, 세션 역할, 감사 로그, 안전 테스트 확인"
    ),
    "기본 흐름: GET /draft -> 분석 생성/조회 -> RAG 검색 -> OpenAI/Bedrock/fallback 초안 생성 -> official_drafts 저장 -> rag_contexts 연결.": (
        "기본 흐름: POST /api/v1/complaints/{id}/draft-runs -> 권위 DB 근거 후보 -> "
        "Python 구조화 초안 -> Spring 결정론 검증 -> 주장·불변 근거 링크 -> 사람 검토."
    ),
    "기본 흐름: health/Swagger 접근 확인, API Key 미포함/포함 요청 검증, audit_logs 저장 확인.": (
        "기본 흐름: health·운영 Swagger 권한 확인, 세션·CSRF·역할 검증, 변경 감사 로그 확인."
    ),
    "기본 흐름: PUT /draft with draftText -> 기존 초안 revise -> draft_revisions에 before/after 저장 -> DraftResponse.status=REVISED 반환.": (
        "기본 흐름: POST /api/v1/drafts/{id}/reviews -> 검토 통과·반려 기록 -> "
        "별도 승인자 결정 대기 또는 DRAFT_REVIEW 복귀."
    ),
    "기본 흐름: 초안 생성 과정에서 PostgresKnowledgeDocumentSearchService가 분석 결과 기반 키워드 검색 -> rag_contexts 저장 -> GET /rag-contexts로 근거 조회.": (
        "기본 흐름: 권위 DB에서 분석 기반 공식 근거 후보 조회 -> Python 구조화 초안 -> "
        "Spring 후보·스키마·근거 검증 -> evidence_snapshots와 claim_evidence_links 저장."
    ),
    "종료조건: complaint_analysis가 저장되고 complaints.status가 ANALYZED로 전환된다.": (
        "종료조건: 검증된 분석과 복합 이슈가 저장되고 complaints.status가 TRIAGE_REVIEW로 전환된다."
    ),
    "종료조건: DraftResponse.status=DRAFT와 draftText, references가 반환된다.": (
        "종료조건: 검증된 주장·근거 링크와 검토용 초안이 저장되고 complaints.status가 DRAFT_REVIEW가 된다."
    ),
    "주의: 개발 기본값은 API Key 인증 비활성화이며, 인증 시험은 설정 활성화 조건부이다.": (
        "주의: 운영은 세션 역할 인증과 CSRF를 필수로 하며 인증 없는 실행은 명시적 dashboard-h2 데모에만 허용한다."
    ),
    "주의: 구현된 API는 POST /analyze가 아니라 GET /analysis이다.": (
        "주의: 분석 요청은 POST /api/v1/complaints/{id}/analysis-runs이며 GET은 상태 조회만 수행한다."
    ),
    "API Key 인증 및 감사 로그": "세션 역할 인증, CSRF 및 변경 감사 로그",
    "API Key 인증 시험은 app.security.api-key.enabled=true 조건에서 수행하는 조건부 보안 시험이다.": (
        "운영 보안 시험은 세션 역할, CSRF, 내부 작업자 서비스 토큰, 변경 감사 로그를 검증한다."
    ),
    "DraftResponse.status=DRAFT, references 1건 이상": (
        "complaints.status=DRAFT_REVIEW, 주장 근거 연결률 100%, 불변 근거 스냅샷 1건 이상"
    ),
    "GET /departments -> 위험물 민원 POST -> GET /analysis -> GET /draft": (
        "위험물 민원 POST -> POST analysis-runs -> 안전 차단 상태와 사람 검토 확인"
    ),
    "GET /draft, GET /rag-contexts": "POST draft-runs, GET /api/v1/complaints/{id}",
    "POST /complaints -> GET /analysis -> GET /draft -> PUT /draft -> GET /rag-contexts": (
        "POST /api/v1/complaints -> POST analysis-runs -> POST draft-runs -> "
        "POST reviews -> POST approvals -> POST complete"
    ),
    "PUT /draft draftText": "POST /api/v1/drafts/{id}/reviews 및 /approvals",
    "민원 접수-분석-RAG-초안 Agent Loop": "민원 접수-분석-근거 검증-검토·승인 작업 흐름",
    "시험 기준은 최종 백엔드 egov-boot-web, Flyway V1~V4, ComplaintApiSmokeTest, 수동 API Key/감사 로그 검증이다.": (
        "시험 기준은 권위 서버 egov-boot-web, Flyway V1~V16, 안전 통합 테스트, "
        "세션·내부 작업자 토큰·감사 로그 검증이다."
    ),
    "위험물 rawText -> GET /analysis": "위험물 비식별 파생본 -> POST analysis-runs -> 안전 차단 검증",
}


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_architecture_document() -> Document:
    document = Document()
    document.add_heading("Evidence-Verified Complaint Support Architecture", 0)
    document.add_paragraph(f"Generated: {DATE}")
    document.add_paragraph(
        "Spring is authoritative for workflow state, permissions, deterministic "
        "verification, audit, human decisions, and job leases. Python workers perform "
        "provider-backed analysis and draft generation plus restricted asynchronous support tasks. "
        "Spring applies a worker result only after governed hash, schema, prompt-version, "
        "and evidence-candidate validation."
    )
    add_table(
        document,
        ["Component", "Responsibility", "Forbidden responsibility"],
        [
            ["egov-boot-web", "State, roles, jobs, evidence, reviews, approvals", "External automatic send"],
            ["Python worker", "AI provider calls, redaction, quarantined attachment extraction, retrieval, verification support", "Approval, completion, or authoritative state changes"],
            ["Internal worker API", "Service-token claim, AI/support result, and failure contract", "Browser access or human workflow decisions"],
            ["Knowledge maintenance", "Official-source freshness and derived purpose-specific OpenSearch indices", "Workflow decisions"],
            ["PostgreSQL", "Versioned state, source registry, immutable evidence snapshots, audit", "None"],
            ["Review workspace", "Explicit human requests and decisions", "Browser API keys"],
        ],
    )
    document.add_paragraph(
        "Draft providers return draft-claims-v1 JSON. Spring validates every claim "
        "and supplied evidence ID before rendering a review draft and linking immutable snapshots. "
        "Official legal evidence requires a source version, legal basis, effective period, "
        "jurisdiction, source URL, verification status, and a content hash that matches the captured text. "
        "Conflicting or incomplete candidates remain visible as non-supporting immutable snapshots."
    )
    document.add_paragraph(
        "Analysis providers return exactly one complaint-support-v1 JSON document. "
        "Untrusted complaint, analysis, and evidence text is boundary-escaped before provider calls."
    )
    document.add_paragraph(
        "Official-law synchronization preserves XML provisions, effective dates, source versions, "
        "and hashes. Stale or failed sources are excluded. Quarantined attachments must pass malware "
        "inspection, extraction, and redaction before their derivatives may enter a provider prompt."
    )
    document.add_paragraph(
        "OpenAI, Bedrock, Mock, and public APIs share a bounded provider runtime with independent "
        "cost limits, retries, and circuit state. Attachment and sensitive-payload objects without "
        "database references are removed only after the configured orphan-cleanup grace period."
    )
    return document


def build_api_document() -> Document:
    document = Document()
    document.add_heading("Versioned Complaint Workflow API", 0)
    document.add_paragraph(
        "Every mutation requires Idempotency-Key. Mutations after intake also "
        "require If-Match with the current entity version. A replay with the "
        "same idempotency key returns the original result before stale-version "
        "validation. GET requests do not mutate workflow or request-audit data."
    )
    add_table(
        document,
        ["Method", "Path", "Purpose"],
        [
            ["POST", "/api/v1/complaints", "Intake a complaint"],
            ["POST", "/api/v1/complaints/{id}/analysis-runs", "Request asynchronous issue analysis"],
            ["POST", "/api/v1/complaints/{id}/draft-runs", "Request evidence-gated draft generation"],
            ["GET", "/api/v1/runs/{id}", "Inspect job status and failure reason"],
            ["POST", "/api/v1/issues/{id}/location-confirmations", "Record human-confirmed location"],
            ["POST", "/api/v1/drafts/{id}/reviews", "Record reviewer decision"],
            ["POST", "/api/v1/drafts/{id}/approvals", "Record separate approver decision"],
            ["POST", "/api/v1/complaints/{id}/complete", "Record manual completion after approval"],
        ],
    )
    document.add_paragraph(
        "Internal Python workers use /internal/v1/worker/jobs/claim, /results, /support-results, "
        "and /failures with a dedicated service token. Spring owns every lease, retry, completion, "
        "failure, and audit transition. These endpoints are not browser APIs."
    )
    return document


def build_test_document() -> Document:
    document = Document()
    document.add_heading("Trust Workflow Verification Scenarios", 0)
    add_table(
        document,
        ["Scenario", "Required result"],
        [
            ["Unverified or missing official evidence", "Draft blocked"],
            ["Non-national or incomplete official source metadata", "Draft blocked before provider call"],
            ["Overlapping official evidence conflicts", "CONFLICT_DETECTED; no draft created"],
            ["Unknown location", "NEEDS_LOCATION until human confirmation"],
            ["Reviewer attempts own approval", "Rejected"],
            ["Stale If-Match version", "Conflict"],
            ["Repeated Idempotency-Key", "Original result returned without duplicate work"],
            ["Concurrent identical Idempotency-Key", "At most one resource created"],
            ["PII in complaint text", "External AI receives redacted derivative"],
            ["Unapproved or malicious attachment", "Analysis and draft blocked"],
            ["Official source sync failure or stale source", "Source excluded from legal evidence"],
            ["OpenSearch hit without governed DB document", "Hit ignored"],
            ["Prompt injection in source material", "Treated as untrusted data"],
            ["Missing or invalid internal worker token", "Worker API request rejected"],
            ["Worker input or output hash changed", "Result rejected; authoritative state unchanged"],
            ["Worker draft selects evidence outside governed candidates", "Result blocked before draft application"],
            ["Retryable provider failure", "FAILED with bounded exponential retry-not-before"],
            ["Expired worker lease", "Late result rejected and job recovered with bounded retry"],
            ["Provider cost or failure threshold exceeded", "Independent provider circuit rejects call"],
            ["Old unreferenced attachment or sensitive payload", "Deleted after orphan-cleanup grace period"],
            ["Untrusted text closes a prompt boundary", "Boundary characters escaped; injected instruction remains data"],
            ["Trailing text after structured AI JSON", "Output blocked by strict single-document parsing"],
            ["Draft claim references unknown evidence ID", "Draft blocked by schema validation"],
            ["Official evidence lacks source version or content hash", "Draft blocked before provider call"],
            ["Official evidence content hash does not match captured text", "Draft blocked; candidate remains visible as non-supporting snapshot"],
            ["HTTP Basic credential sent to production API", "Not authenticated; form-login session required"],
            ["Approved complaint completion", "Manual record only; no external send"],
            ["150-case synthetic golden evaluation", "Recall@10 and department Top-3 >= 0.95; claim coverage 100%; safety failures 0"],
        ],
    )
    return document


def save(document: Document, filename: str) -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUT_DIR / filename)


def iter_paragraphs(parent):
    yield from parent.paragraphs
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def update_legacy_corrected_documents() -> None:
    for path in OUT_DIR.glob("*수정본.docx"):
        document = Document(path)
        changed = False
        for paragraph in iter_paragraphs(document):
            replacement = LEGACY_CORRECTED_REPLACEMENTS.get(paragraph.text)
            if replacement is not None:
                paragraph.text = replacement
                changed = True
        if changed:
            document.save(path)


def main() -> None:
    save(build_architecture_document(), "trust-workflow-architecture.docx")
    save(build_api_document(), "trust-workflow-api.docx")
    save(build_test_document(), "trust-workflow-tests.docx")
    update_legacy_corrected_documents()


if __name__ == "__main__":
    main()
